from docx import Document, enum
from sys import stdin

doc = Document()
place = input()
time = input()
p = 2
for fio in stdin:
    doc.add_heading('Приглашение', 0)
    doc.add_heading('Дорогая ' + fio + '!', 1)
    doc.add_paragraph('Приглашаем тебя на празднование дня 8 марта,' +
                      ' которое состоится ' + time + ' ' + place)
    doc.paragraphs[p].runs[0].add_break(enum.text.WD_BREAK.PAGE)
    p += 3
doc.save('invitations.docx')