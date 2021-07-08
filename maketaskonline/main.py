import os
import re
import uuid
from random import choice

from flask import Flask, render_template, make_response, session, request, url_for
from flask_bootstrap import Bootstrap
from flask_login import LoginManager, login_user, login_required, logout_user, current_user
from werkzeug.exceptions import abort
from werkzeug.utils import redirect

from data import db_session
from data.section import Section
from data.subjects import Subjects
from data.users import User
from data.task import Task
from data.variables import Variables
from data.worktasklist import Worktasklist
from data.taskset import TaskSet
from data.tasksetline import TaskSetLine
from data.variants import Variants
from forms.user import RegisterForm, LoginForm, SubjectForm, SectionForm, TaskForm, VariablesForm, ToWorkForm, \
    TaskListForm, VariantsForm
from docx import Document

app = Flask(__name__)
app.config['SECRET_KEY'] = 'HSa6jK1Rb0zMDEPoTlvf5SYg4I6FtkaK'
bootstrap = Bootstrap(app)
login_manager = LoginManager()
login_manager.init_app(app)


# Проверка доступности объекта текущему пользователю
def available(object, object_id):
    db_sess = db_session.create_session()
    if object == "subjects":
        subjects = db_sess.query(Subjects).filter(Subjects.id == object_id).one()
        if not subjects:
            return False
        if subjects.user_id == int(current_user.get_id()):
            return True
    elif object == "taskset":
        taskset = db_sess.query(TaskSet).filter(TaskSet.id == object_id).one()
        if not taskset:
            return False
        if taskset.user_id == int(current_user.get_id()):
            return True
    elif object == "variants":
        variants = db_sess.query(Variants).filter(Variants.id == object_id).one()
        if variants:
            return available('taskset', variants.taskset_id)
    elif object == "section":
        section = db_sess.query(Section).filter(Section.id == object_id).one()
        if section:
            return available('subjects', section.subject_id)
    elif object == "task":
        task = db_sess.query(Task).filter(Task.id == object_id).one()
        if task:
            return available('section', task.section_id)
    elif object == "variables":
        variable = db_sess.query(Variables).filter(Variables.id == object_id).one()
        if variable:
            return available('task', variable.task_id)
    return False


@login_manager.user_loader
def load_user(user_id):
    db_sess = db_session.create_session()
    return db_sess.query(User).get(user_id)


@app.route("/")
@app.route("/index")
def index():
    db_sess = db_session.create_session()
    subjects = []
    if current_user.is_authenticated:
        subjects = db_sess.query(Subjects).filter(
            (Subjects.user_id == current_user.get_id()) | (not Subjects.is_private))
    return render_template('index.html', subjects=subjects)


@app.route("/worktasklist", methods=['GET', 'POST'])
@login_required
def worktasklist():
    db_sess = db_session.create_session()
    worktask = db_sess.query(Worktasklist).filter(Worktasklist.user_id == current_user.get_id()).all()
    form = TaskListForm()
    if form.validate_on_submit():
        db_sess = db_session.create_session()
        worktasklist = db_sess.query(Worktasklist).filter(Worktasklist.user_id == current_user.get_id())
        taskset = TaskSet(
            user_id=current_user.get_id(),
            name=form.name.data
        )
        db_sess.add(taskset)
        db_sess.commit()
        db_sess.flush()
        for task in worktasklist:
            tl = TaskSetLine(
                taskset_id=taskset.id,
                task_id=task.task_id,
                amount=task.amount
            )
            db_sess.add(tl)
        worktasklist = db_sess.query(Worktasklist).filter(Worktasklist.user_id == current_user.get_id()).delete()
        db_sess.commit()
        return redirect(url_for('taskset'))
    return render_template('worktasklist.html', worktask=worktask, form=form, title="Список задач в работе")

@app.route("/taskset")
@login_required
def taskset():
    db_sess = db_session.create_session()
    taskset = db_sess.query(TaskSet).filter(TaskSet.user_id == current_user.get_id())
    return render_template('taskset.html', taskset=taskset, title="Список работ")


@app.route("/variants/<taskset_id>", methods=['GET', 'POST'])
@login_required
def variants(taskset_id):
    form = VariantsForm()
    if not available("taskset", taskset_id):
        return redirect("index")
    db_sess = db_session.create_session()
    variants = db_sess.query(Variants).filter(Variants.taskset_id == taskset_id)
    if form.validate_on_submit():
        variant_add(taskset_id, form.amount.data)
        variants = db_sess.query(Variants).filter(Variants.taskset_id == taskset_id)
        return render_template('variants.html', variants=variants, form=form, title="Список вариантов")
    return render_template('variants.html', variants=variants, form=form, title="Список вариантов")


def generate_variant_and_key(taskset_id, number):
    db_sess = db_session.create_session()
    tasksetline = db_sess.query(TaskSetLine).filter(TaskSetLine.taskset_id == taskset_id)
    num_task = 1
    text_variant = "\nВариант №" + str(number)
    text_key = "Ключ к варианту №" + str(number)
    for row in tasksetline:
        task_id = row.task_id
        amount = row.amount
        for _ in range(amount):
            condition, answer = generate_condition_and_answer(num_task, task_id)
            num_task += 1
            text_variant += "\n" + condition
            text_key += "\n" + answer
    return text_variant, text_key


def generate_condition_and_answer(number, task_id):
    text_condition = "Задание № " + str(number)
    db_sess = db_session.create_session()
    task = db_sess.query(Task).filter(Task.id == task_id).one()
    variables = db_sess.query(Variables).filter(Variables.task_id == task_id)
    condition = task.condition
    formula = task.formula
    for var in variables:
        var_name = '{' + var.name + '}'
        if "range" in var.range:
            var_value = str(choice(list(eval(var.range))))
        elif var.range[0] == '[' and var.range[-1] == ']':
            var_value = str(choice(eval(var.range)))
        else:
            var_value = 'не определен диапазон'
        condition = condition.replace(var_name, var_value)
        formula = formula.replace(var_name, var_value)
    text_condition += "\n" + condition
    text_answer = str(number) + ")" + str(eval(formula))
    return text_condition, text_answer


def variant_add(taskset_id, amount):
    db_sess = db_session.create_session()
    content = ""
    answers = ""
    for i in range(1, amount + 1):
        variant, answer = generate_variant_and_key(taskset_id, i)
        content += variant + "\n"
        answers += answer + "\n"
    file_name = 'file' + str(uuid.uuid1()) + '.docx'
    full_file_name = os.path.join(str(os.getcwd()), 'static', 'output', file_name)
    document = Document()
    document.add_paragraph(content)
    document.add_paragraph(answers)
    document.save(full_file_name)
    variants = Variants(
        taskset_id=taskset_id,
        content=content,
        answers=answers,
        file=url_for('static', filename='output/' + file_name)
    )
    db_sess.add(variants)
    db_sess.commit()
    return redirect(url_for('variants', taskset_id=taskset_id))



@app.route("/delworktask/<worktask_id>")
@login_required
def delworktask(worktask_id):
    db_sess = db_session.create_session()
    worktask = db_sess.query(Worktasklist).filter(Worktasklist.id == worktask_id).one()
    db_sess.delete(worktask)
    db_sess.commit()
    return redirect(url_for("worktasklist"))


@app.route("/section/<subject_id>")
@login_required
def section(subject_id):
    if not available("subjects", subject_id):
        return redirect("index")
    db_sess = db_session.create_session()
    subject = db_sess.query(Subjects).filter(Subjects.id == subject_id).one()
    sections = db_sess.query(Section).filter(Section.subject_id == subject_id)
    return render_template('section.html', sections=sections, subject=subject)


@app.route("/task/<subject_id>/<section_id>")
@login_required
def task(subject_id, section_id):
    if not(available("subjects", subject_id) and available("section", section_id)):
        return redirect("index")
    db_sess = db_session.create_session()
    subject = db_sess.query(Subjects).filter(Subjects.id == subject_id).one()
    section = db_sess.query(Section).filter(Section.id == section_id).one()
    tasks = db_sess.query(Task).filter(Task.section_id == section_id)
    return render_template('task.html', section=section, subject=subject, tasks=tasks)


@app.route("/variables/<subject_id>/<section_id>/<task_id>", methods=['GET', 'POST'])
@login_required
def variables(subject_id, section_id, task_id):
    if not(available("subjects", subject_id) and
           available("section", section_id) and (available("task", task_id))):
        return redirect("index")
    db_sess = db_session.create_session()
    subject = db_sess.query(Subjects).filter(Subjects.id == subject_id).one()
    section = db_sess.query(Section).filter(Section.id == section_id).one()
    task = db_sess.query(Task).filter(Task.id == task_id).one()
    variables = db_sess.query(Variables).filter(Variables.task_id == task_id).all()
    if not variables:
        vars = sorted(list(set(re.findall(r"\{(.*?)\}", task.condition))))
        db_sess = db_session.create_session()
        for var in vars:
            variable = Variables(
                name=var,
                task_id=task.id,
                type="",
                range=""
            )
            db_sess.add(variable)
        db_sess.commit()
    variables = db_sess.query(Variables).filter(Variables.task_id == task_id)
    form = ToWorkForm()
    if form.validate_on_submit():
        db_sess = db_session.create_session()
        wortasklist = Worktasklist(
            task_id=task_id,
            user_id=current_user.get_id(),
            amount=form.amount.data
        )
        db_sess.add(wortasklist)
        db_sess.commit()
        return redirect(url_for('worktasklist'))

    return render_template('variables.html', section=section, subject=subject, task=task,
                           variables=variables, form=form)


@app.route("/edit_var/<subject_id>/<section_id>/<task_id>/<var_id>", methods=["GET", "POST"])
@login_required
def edit_var(subject_id, section_id, task_id, var_id):
    if not(available("subjects", subject_id) and available("section", section_id) and
           (available("task", task_id)) and available("variables", var_id)):
        return redirect("index")
    form = VariablesForm()
    if request.method == "GET":
        db_sess = db_session.create_session()
        var = db_sess.query(Variables).filter(Variables.id == var_id).one()
        if var:
            form.type.data = var.type
            form.range.data = var.range
        else:
            abort(404)
    if form.validate_on_submit():
        db_sess = db_session.create_session()
        var = db_sess.query(Variables).filter(Variables.id == var_id).one()
        if var:
            var.type = form.type.data
            var.range = form.range.data
            db_sess.commit()
            return redirect(url_for('variables', subject_id=subject_id, section_id=section_id, task_id=task_id))
            # return render_template('variables.html', section=section, subject=subject, task=task, variables=variables)
        else:
            abort(404)
    return render_template('edit_var.html', title="Редактирование переменной", form=form, var=var)


@app.route('/register', methods=['GET', 'POST'])
def register():
    form = RegisterForm()
    if form.validate_on_submit():
        if form.password.data != form.password_again.data:
            return render_template('register.html', title='Регистрация',
                                   form=form,
                                   message="Пароли не совпадают")
        db_sess = db_session.create_session()
        if db_sess.query(User).filter(User.email == form.email.data).first():
            return render_template('register.html', title='Регистрация',
                                   form=form,
                                   message="Такой пользователь уже есть")
        user = User(
            name=form.name.data,
            email=form.email.data
        )
        user.set_password(form.password.data)
        db_sess.add(user)
        db_sess.commit()
        return redirect('/login')
    return render_template('register.html', title='Регистрация', form=form)


@login_required
@app.route('/subject_add', methods=['GET', 'POST'])
def subject_add():
    form = SubjectForm()
    if form.validate_on_submit():
        db_sess = db_session.create_session()
        subject = Subjects(
            name=form.name.data,
            is_private=form.is_private.data,
            user_id=current_user.get_id()
        )
        db_sess.add(subject)
        db_sess.commit()
        return redirect('/')
    return render_template('form_add.html', title='Добавление предмета', form=form)


@app.route('/section_add/<subject_id>', methods=['GET', 'POST'])
@login_required
def section_add(subject_id):
    if not available("subjects", subject_id):
        return redirect("index")
    form = SectionForm()
    if form.validate_on_submit():
        db_sess = db_session.create_session()
        section = Section(
            name=form.name.data,
            subject_id=subject_id
        )
        db_sess.add(section)
        db_sess.commit()
        return redirect(f"/section/{subject_id}")
    return render_template('form_add.html', title='Добавление раздела', form=form)


@app.route('/task_add/<subject_id>/<section_id>', methods=['GET', 'POST'])
@login_required
def task_add(subject_id, section_id):
    if not(available("subjects", subject_id) and available("section", section_id)):
        return redirect("index")
    form = TaskForm()
    if form.validate_on_submit():
        db_sess = db_session.create_session()
        task = Task(
            name=form.name.data,
            section_id=section_id,
            condition=form.condition.data,
            formula=form.formula.data
        )
        db_sess.add(task)
        db_sess.commit()
        return redirect(f"/task/{subject_id}/{section_id}")
    return render_template('form_add.html', title='Добавление задачи', form=form)


@app.route("/login", methods=["POST", "GET"])
def login():
    form = LoginForm()
    if form.validate_on_submit():
        db_sess = db_session.create_session()
        user = db_sess.query(User).filter(User.email == form.email.data).first()
        if user and user.check_password(form.password.data):
            login_user(user, remember=form.remember_me.data)
            return redirect("/")
        return render_template('login.html',
                               message="Неправильный логин или пароль",
                               form=form)

    return render_template('login.html', title='Вход', form=form)


@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect("/")


def main():
    db_session.global_init("db/maketask.sqlite")
    app.run('127.0.0.1', 8080, debug=True)


if __name__ == '__main__':
    main()
